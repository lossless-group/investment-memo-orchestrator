#!/usr/bin/env python3
"""
Generate Hypernova-branded Word reference document for pandoc.

This creates templates/hypernova-reference.docx with branded styles
that pandoc uses when converting markdown to Word format.

Usage:
    python create-word-reference.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# Hypernova Brand Colors (from hypernova-style.css)
COLORS = {
    'navy': RGBColor(26, 58, 82),      # #1a3a52
    'cyan': RGBColor(29, 211, 211),    # #1dd3d3
    'dark_gray': RGBColor(26, 35, 50), # #1a2332
    'light_gray': RGBColor(107, 114, 128), # #6b7280
    'white': RGBColor(255, 255, 255),  # #ffffff
    'cream': RGBColor(240, 240, 235),  # #f0f0eb
}


def set_cell_border(cell, **kwargs):
    """
    Set cell borders for table cells.

    Args:
        cell: docx table cell
        **kwargs: top, bottom, left, right, insideH, insideV with dict values
                  containing 'sz' (size), 'val' (style), 'color' (hex without #)
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)


def create_reference_document(output_path: Path):
    """
    Create a Word reference document with Hypernova branding.

    Args:
        output_path: Path where to save the reference document
    """

    doc = Document()

    # Set up page layout
    section = doc.sections[0]
    section.page_height = Inches(11)  # Letter size
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # Configure styles
    styles = doc.styles

    # Normal (body text) style
    normal = styles['Normal']
    normal_font = normal.font
    normal_font.name = 'Calibri'  # Fallback (Arboria if user has it)
    normal_font.size = Pt(11)
    normal_font.color.rgb = COLORS['dark_gray']

    normal_paragraph = normal.paragraph_format
    normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_paragraph.line_spacing = 1.5
    normal_paragraph.space_after = Pt(6)

    # Heading 1 style
    h1 = styles['Heading 1']
    h1_font = h1.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(20)
    h1_font.bold = True
    h1_font.color.rgb = COLORS['navy']

    h1_paragraph = h1.paragraph_format
    h1_paragraph.space_before = Pt(24)
    h1_paragraph.space_after = Pt(12)

    # Heading 2 style
    h2 = styles['Heading 2']
    h2_font = h2.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(16)
    h2_font.bold = True
    h2_font.color.rgb = COLORS['navy']

    h2_paragraph = h2.paragraph_format
    h2_paragraph.space_before = Pt(18)
    h2_paragraph.space_after = Pt(10)

    # Heading 3 style
    h3 = styles['Heading 3']
    h3_font = h3.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_font.color.rgb = COLORS['navy']

    h3_paragraph = h3.paragraph_format
    h3_paragraph.space_before = Pt(12)
    h3_paragraph.space_after = Pt(8)

    # Heading 4 style
    h4 = styles['Heading 4']
    h4_font = h4.font
    h4_font.name = 'Calibri'
    h4_font.size = Pt(12)
    h4_font.bold = True
    h4_font.color.rgb = COLORS['dark_gray']

    h4_paragraph = h4.paragraph_format
    h4_paragraph.space_before = Pt(10)
    h4_paragraph.space_after = Pt(6)

    # Hyperlink style
    try:
        hyperlink = styles['Hyperlink']
        hyperlink_font = hyperlink.font
        hyperlink_font.color.rgb = COLORS['cyan']
    except KeyError:
        pass  # Hyperlink style doesn't exist in all templates

    # Add header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header left: Company name
    run_left = header_para.add_run("Hypernova Capital")
    run_left.font.name = 'Calibri'
    run_left.font.size = Pt(10)
    run_left.font.bold = True
    run_left.font.color.rgb = COLORS['navy']

    # Add tab for right-aligned title
    header_para.add_run("\t")

    # Header right: Document title (placeholder)
    run_right = header_para.add_run("Investment Memo")
    run_right.font.name = 'Calibri'
    run_right.font.size = Pt(10)
    run_right.font.color.rgb = COLORS['light_gray']

    # Set tab stop for right alignment
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    # Add footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer: Confidential | Page # | Tagline
    run_conf = footer_para.add_run("Confidential")
    run_conf.font.name = 'Calibri'
    run_conf.font.size = Pt(9)
    run_conf.font.italic = True
    run_conf.font.color.rgb = COLORS['light_gray']

    footer_para.add_run(" | ")

    # Page number
    run_page = footer_para.add_run()
    run_page.font.name = 'Calibri'
    run_page.font.size = Pt(9)
    run_page.font.color.rgb = COLORS['light_gray']

    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_page._element.append(fldChar1)
    run_page._element.append(instrText)
    run_page._element.append(fldChar2)

    footer_para.add_run(" | ")

    run_tag = footer_para.add_run("Network-Driven | High-impact | Transformative")
    run_tag.font.name = 'Calibri'
    run_tag.font.size = Pt(9)
    run_tag.font.color.rgb = COLORS['light_gray']

    # Add sample content to establish styles (will be removed)
    doc.add_heading('Sample Heading 1', level=1)
    doc.add_paragraph('This is sample body text to establish the Normal style.')

    doc.add_heading('Sample Heading 2', level=2)
    p = doc.add_paragraph('Bullet point one')
    p.style = 'List Bullet'

    doc.add_heading('Sample Heading 3', level=3)

    # Add sample table to establish table style
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    # Style header row
    header_cells = table.rows[0].cells
    for cell in header_cells:
        cell.text = 'Header'
        # Set background to navy
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1a3a52')  # Navy
        cell._element.get_or_add_tcPr().append(shading_elm)

        # Set text to white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = COLORS['white']
                run.font.bold = True

    # Add data to body rows
    for i in range(1, 3):
        cells = table.rows[i].cells
        for j, cell in enumerate(cells):
            cell.text = f'Data {i}-{j+1}'

    # Save the document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"✓ Created Word reference document: {output_path}")
    print(f"  - Hypernova brand colors applied")
    print(f"  - Heading styles (H1-H4) configured")
    print(f"  - Header: 'Hypernova Capital' | Document title")
    print(f"  - Footer: 'Confidential' | Page # | Tagline")
    print(f"  - Body text: 11pt Calibri, justified, dark gray")
    print(f"\nUse with:")
    print(f"  python md2docx.py memo.md --reference-doc {output_path}")


def main():
    output_path = Path(__file__).parent / "templates" / "hypernova-reference.docx"

    print("Generating Hypernova-branded Word reference document...")
    print()

    try:
        create_reference_document(output_path)
    except Exception as e:
        print(f"✗ Error creating reference document: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == '__main__':
    import sys
    sys.exit(main())
